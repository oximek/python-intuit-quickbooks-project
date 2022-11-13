from docx import Document
from docx.shared import Inches


document = Document()

#profile picture
document.add_picture(
    'me.jpg.jpg', 
    width=Inches(2.0)
)

#name phone number and emails details
name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

#about me
document.add_heading('About me')
about_me = input('Tell about yourself? ')
document.add_paragraph(about_me)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
        'Describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break
            

document.save('cv.docx')